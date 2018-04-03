#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys
from docx_simple_service import SimpleDocxService
from docx.shared import Cm, Inches, Pt
import json
import os
import glob

docx = SimpleDocxService()


def open_report_json(fname: str):
    fp = open(fname, 'r')
    report = json.load(fp)
    return report


def build_report(report):
    """
    レポートを組み立てます

    :return:
    """
    for case in report["results"]:
        # 文節タイトル表示
        docx.add_head(case["name"], 2)
        with docx.open_text() as text:
            evaluation = []
            text.add("日付: %s\n" % case["date"])
            text.add("URL: %s\n" % case["url"])
            text.add("テスト結果:")
            if case["result"] == "failed":
                text.add("%s\n" % case["result"]).color(*RED)
            else:
                text.add("%s\n" % case["result"]).color(0, 128, 0)
            request = case["request"]
            response = case["response"]
            evaluation = case["evaluation"]
            method = request["method"]
            query = case["query"]

            text.add("メソッド: %s\n" % method)
            docx.add_head("REQUEST", 3)
            docx.add_head("ヘッダ", 4)
            with docx.open_text() as text:
                length = len(request["header"])
                table = docx.add_table(row=length, col=2)
                table.autofit = False
                table.style = 'Table Grid'
                set_col_widths(table)
                request_header(request, table)
                set_table_format(table)

            docx.add_head("クエリ", 4)
            with docx.open_text() as text:
                length = len(query)
                table = docx.add_table(row=length, col=2)
                table.autofit = False
                table.style = 'Table Grid'
                set_col_widths(table)
                query_param(query, table)
                set_table_format(table)

            docx.add_head("RESPONSE", 3)
            docx.add_head("ヘッダ", 4)
#
            with docx.open_text() as res:
                length = len(response["header"])
                table = docx.add_table(row=length, col=2)
                table.autofit = False
                table.style = 'Table Grid'
                set_col_widths(table)
                response_header(response, table)
                set_table_format(table)

            docx.add_head("Payload", 4)
            with docx.open_text() as res_body:
                res_body.add("%s" % response["payload"])
            #
            docx.add_head("評価結果", 3)
            docx.add_ordered_list(evaluation)


def query_param(query, table):
    for i, q in enumerate(query):
        query_string = str(query[q])
        if len(query_string) > 512:
            value = query[q]
            head = value[0:30]
            tail = value[-30:]
            v = "%s (snip) %s" % (head, tail)
            cell = table.cell(i, 0)
            cell.text = q
            cell = table.cell(i, 1)
            cell.text = v

        else:
            cell = table.cell(i, 0)
            cell.text = q
            cell = table.cell(i, 1)
            print(query[q])
            cell.text = query_string


def request_header(request, table):
    for i, header in enumerate(request["header"]):
        if len(request["header"][header]) > 512:
            value = request["header"][header]
            head = value[0:30]
            tail = value[-30:]
            v = "%s (snip) %s" % (head, tail)
            cell = table.cell(i, 0)
            cell.text = header
            cell = table.cell(i, 1)
            cell.text = v

        else:
            cell = table.cell(i, 0)
            cell.text = header
            cell = table.cell(i, 1)
            cell.text = request["header"][header]

def response_header(response, table):

    for i, val in enumerate(response["header"]):
        print(response["header"][val])
        str_val = str(response["header"][val])
        if len(str_val) > 512:
            value = response["header"][val]
            head = value[0:30]
            tail = value[-30:]
            v = "%s (snip) %s" % (head, tail)
            cell = table.cell(i, 0)
            cell.text = val
            cell = table.cell(i, 1)
            cell.text = v
        #                        text.add("%s:%s\n" % (header, v))
        else:
            cell = table.cell(i, 0)
            cell.text = str(val)
            cell = table.cell(i, 1)
            cell.text = str(response["header"][val])

def set_table_format(table):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)
                    font.name = "Courier New"
#                        text.add("%s:%s\n" % (header, request["header"][header]))

def set_col_widths(table):
    widths = (Inches(2.0), Inches(4))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

if __name__ == "__main__":

    RED = (0xff, 0x00, 0x00)  # (Red, Green, Blue)
    GREEN = (0x00, 0xFF, 0x00)
    #フォント設定
    docx.set_normal_font("Courier New", 9)

    # タイトル表示
    docx.add_head(u"テスト結果", 0)

    # 挿絵挿入
    docx.add_picture("report_top.jpg", 3.0)

    reports = glob.glob("result_*")

    for f_name in reports:
        report = open_report_json(f_name)
        docx.add_head(f_name, 1)
        build_report(report)
        docx.add_page_break()

    # セーブです。
    docx.save("test.docx")

    print("complete.")
