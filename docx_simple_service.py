#!/usr/bin/env python
# -*- coding: utf-8 -*-

#
# SimpleDocService
#  python-doxに関する簡単なサービスを提供します。
#  まぁ docxライブラリを理解するためのコードですね。
#

from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
import unicodedata
import re
class SimpleDocxService:

    def __init__(self):
        self.document = Document()

    def set_normal_font(self, name, size):
        """フォントの設定"""
        font = self.document.styles['Normal'].font
        font.name = name
        font.size = Pt(size)

    def add_head(self, text, lv):
        """見出しの設定"""
        self.document.add_heading(text, level=lv)

    def open_text(self):
        """テキスト追加開始"""
        class Paragraph:

            def __init__(self, paragraph):
                self.paragraph = paragraph
                self.text = None

            def __enter__(self):
                """テキスト追加開始"""
                return self

            def add(self, text, encode=None):
                """テキスト追加"""
#                if encode:
#                    text = unicode(text, encode)
                self.text = self.paragraph.add_run(text)
                return self

            def italic(self):
                """斜体にする"""
                self.text.italic = True
                return self

            def bold(self):
                """太字にする"""
                self.text.bold = True
                return self

            def color(self, r, g, b):
                """色を付ける"""
                self.text.font.color.rgb = RGBColor(r, g, b)
                return self

            def close(self):
                """テキスト追加終了"""
                del self.paragraph
                del self.text

            def __exit__(self, *args):
                """テキスト追加終了"""
                self.close()

        return Paragraph(self.document.add_paragraph())

    def add_ordered_list(self, text):
        for line in text:
            if re.search(r"Fail!$", line):
                self.paragraph =  self.document.add_paragraph(line, style='ListBullet')
            else:
                self.paragraph = self.document.add_paragraph(line, style='ListBullet')

    def add_picture(self, filename, inch):
        """図を挿入する"""
        self.document.add_picture(filename, width=Inches(inch))

    def add_page_break(self):
        """改ページする"""
        self.document.add_page_break()

    def add_table(self, row, col):
        table = self.document.add_table(row, col)
        return table

    def save(self, name):
        """docxファイルとして出力"""
        self.document.save(name)
